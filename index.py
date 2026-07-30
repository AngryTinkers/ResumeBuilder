import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


RESUME_DIR = Path(__file__).resolve().parent / "resume"
TEMPLATE_PATH = RESUME_DIR / "template.docx"
CONTENT_PATH = RESUME_DIR / "content.txt"
OUTPUT_PATH = RESUME_DIR / "Kevin Grant.docx"


def parse_bold_segments(s: str) -> list[tuple[bool, str]]:
  """Split *s* into (bold, text) segments using ``**...**`` pairs."""
  if not s:
    return []
  out: list[tuple[bool, str]] = []
  pattern = re.compile(r"\*\*(.+?)\*\*")
  last = 0
  for m in pattern.finditer(s):
    if m.start() > last:
      out.append((False, s[last:m.start()]))
    out.append((True, m.group(1)))
    last = m.end()
  if last < len(s):
    out.append((False, s[last:]))
  return [(b, t) for b, t in out if t]


def merge_adjacent_segments(
  segments: list[tuple[bool, str]],
) -> list[tuple[bool, str]]:
  if not segments:
    return []
  cur_b, cur_t = segments[0]
  merged: list[tuple[bool, str]] = []
  for b, t in segments[1:]:
    if b == cur_b:
      cur_t += t
    else:
      if cur_t:
        merged.append((cur_b, cur_t))
      cur_b, cur_t = b, t
  if cur_t:
    merged.append((cur_b, cur_t))
  return merged


def _first_run_rpr_element(p_el):
  for child in p_el:
    if child.tag == qn("w:r"):
      rpr = child.find(qn("w:rPr"))
      return deepcopy(rpr) if rpr is not None else OxmlElement("w:rPr")
  return OxmlElement("w:rPr")


def _rpr_for_segment(base_rpr, bold: bool):
  rpr = deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")
  b_el = rpr.find(qn("w:b"))
  if bold:
    if b_el is None:
      rpr.append(OxmlElement("w:b"))
  elif b_el is not None:
    rpr.remove(b_el)
  return rpr


def replace_paragraph_formatted(
  paragraph, segments: list[tuple[bool, str]]
) -> None:
  """Replace all runs in *paragraph* with *segments*, keeping paragraph properties."""
  merged = merge_adjacent_segments(segments)
  p_el = paragraph._element
  base_rpr = _first_run_rpr_element(p_el)

  for child in list(p_el):
    if child.tag == qn("w:r"):
      p_el.remove(child)

  for bold, chunk in merged:
    r = OxmlElement("w:r")
    r.append(_rpr_for_segment(base_rpr, bold))
    t = OxmlElement("w:t")
    if chunk.startswith(" ") or chunk.endswith(" ") or "\n" in chunk:
      t.set(qn("xml:space"), "preserve")
    t.text = chunk
    r.append(t)
    p_el.append(r)


def replace_placeholder_formatted(paragraph, placeholder: str, source: str) -> bool:
  """Replace *placeholder* in *paragraph* with *source*, honoring ``**...**`` as bold."""
  if paragraph is None or placeholder not in paragraph.text:
    return False
  before, _, after = paragraph.text.partition(placeholder)
  segments: list[tuple[bool, str]] = []
  segments.extend(parse_bold_segments(before))
  segments.extend(parse_bold_segments(source))
  segments.extend(parse_bold_segments(after))
  replace_paragraph_formatted(paragraph, segments)
  return True


def parse_content(text: str) -> dict:
  # Summary runs until the next ``## `` section heading (e.g. ``## Exp1``).
  summary_match = re.search(
    r"## Summary\s*(.*?)(?=\n## |\Z)", text, re.DOTALL | re.IGNORECASE
  )
  summary = summary_match.group(1).strip() if summary_match else ""
  summary = " ".join(summary.split())

  experiences = []
  for n in range(1, 5):
    exp_match = re.search(
      rf"## Exp{n}\s*(.*?)(?=\n## |\Z)", text, re.DOTALL | re.IGNORECASE
    )
    chunk = exp_match.group(1).strip() if exp_match else ""
    experiences.append(_parse_bullets(chunk))

  tech_match = re.search(r"## Skills\s*(.*)\Z", text, re.DOTALL | re.IGNORECASE)
  tech_raw = tech_match.group(1).strip() if tech_match else ""
  tech_lines = []
  for line in tech_raw.splitlines():
    line = line.strip()
    if line:
      tech_lines.append(line)

  return {"summary": summary, "experiences": experiences, "tech_lines": tech_lines}


def _parse_bullets(chunk: str) -> list[str]:
  bullets = []
  for line in chunk.strip().splitlines():
    stripped = line.strip()
    if stripped.startswith("- "):
      bullets.append(stripped[2:].strip())
  return bullets


def iter_all_paragraphs(doc: Document):
  for p in doc.paragraphs:
    yield p
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for p in cell.paragraphs:
          yield p


def find_paragraph(doc: Document, needle: str):
  for p in iter_all_paragraphs(doc):
    if needle in p.text:
      return p
  return None


def insert_extra_list_items_after(bullet2_paragraph, extra_texts: list[str]) -> None:
  if not extra_texts:
    return

  ref = bullet2_paragraph._element
  template_el = deepcopy(ref)
  parent = bullet2_paragraph._parent
  for extra in extra_texts:
    new_el = deepcopy(template_el)
    new_p = Paragraph(new_el, parent)
    replace_paragraph_formatted(new_p, parse_bold_segments(extra))
    ref.addnext(new_el)
    ref = new_el


def build_resume() -> None:
  raw = CONTENT_PATH.read_text(encoding="utf-8")
  data = parse_content(raw)

  shutil.copyfile(TEMPLATE_PATH, OUTPUT_PATH)
  doc = Document(OUTPUT_PATH)

  replace_placeholder_formatted(
    find_paragraph(doc, "{{ summary }}"), "{{ summary }}", data["summary"]
  )

  exp_keys = ["exp1", "exp2", "exp3", "exp4"]
  for i, key in enumerate(exp_keys):
    bullets = data["experiences"][i] if i < len(data["experiences"]) else []
    b1 = bullets[0] if bullets else ""
    b2 = bullets[1] if len(bullets) > 1 else ""
    extras = bullets[2:]

    p_b1 = find_paragraph(doc, f"{{{{ {key}_bullet1 }}}}")
    p_b2 = find_paragraph(doc, f"{{{{ {key}_bullet2 }}}}")

    replace_placeholder_formatted(p_b1, f"{{{{ {key}_bullet1 }}}}", b1)
    replace_placeholder_formatted(p_b2, f"{{{{ {key}_bullet2 }}}}", b2)

    if p_b2 and extras:
      insert_extra_list_items_after(p_b2, extras)

  tech = data["tech_lines"]
  skill1 = tech[0] if len(tech) > 0 else ""
  skill2 = tech[1] if len(tech) > 1 else ""

  p_s1 = find_paragraph(doc, "{{ skill_cat1 }}")
  p_s2 = find_paragraph(doc, "{{ skill_cat2 }}")

  replace_placeholder_formatted(p_s1, "{{ skill_cat1 }}", skill1)
  replace_placeholder_formatted(p_s2, "{{ skill_cat2 }}", skill2)

  extra_skills = tech[2:] if len(tech) > 2 else []
  if p_s2 and extra_skills:
    insert_extra_list_items_after(p_s2, extra_skills)

  doc.save(OUTPUT_PATH)


if __name__ == "__main__":
  build_resume()
  print(f"Wrote {OUTPUT_PATH}")
